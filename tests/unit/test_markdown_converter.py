"""Unit tests for the markdown-to-DOCX converter module.

Tests cover: tables, blockquotes, links, nested lists, and existing features
(headings, bold, italic, code, flat lists).
"""

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

from patent_system.export.markdown_converter import add_markdown_to_document


# ---------------------------------------------------------------------------
# 8.13 — Table conversion
# ---------------------------------------------------------------------------


class TestTableConversion:
    """Markdown table converts to DOCX table with correct rows/columns."""

    def test_simple_table_structure(self) -> None:
        doc = Document()
        md = "| A | B | C |\n|---|---|---|\n| 1 | 2 | 3 |\n| 4 | 5 | 6 |"
        add_markdown_to_document(doc, md)

        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.rows) == 3  # 1 header + 2 body rows
        assert len(table.columns) == 3

    def test_table_header_content(self) -> None:
        doc = Document()
        md = "| Name | Age |\n|------|-----|\n| Alice | 30 |"
        add_markdown_to_document(doc, md)

        table = doc.tables[0]
        header_cells = [cell.text for cell in table.rows[0].cells]
        assert header_cells == ["Name", "Age"]

    def test_table_body_content(self) -> None:
        doc = Document()
        md = "| X | Y |\n|---|---|\n| a | b |\n| c | d |"
        add_markdown_to_document(doc, md)

        table = doc.tables[0]
        row1 = [cell.text for cell in table.rows[1].cells]
        row2 = [cell.text for cell in table.rows[2].cells]
        assert row1 == ["a", "b"]
        assert row2 == ["c", "d"]

    def test_table_has_borders(self) -> None:
        doc = Document()
        md = "| H1 | H2 |\n|---|---|\n| v1 | v2 |"
        add_markdown_to_document(doc, md)

        table = doc.tables[0]
        tbl_pr = table._tbl.tblPr
        borders = tbl_pr.findall(qn("w:tblBorders"))
        assert len(borders) >= 1


# ---------------------------------------------------------------------------
# 8.14 — Blockquote conversion
# ---------------------------------------------------------------------------


class TestBlockquoteConversion:
    """Markdown blockquote converts to indented paragraph."""

    def test_blockquote_text_preserved(self) -> None:
        doc = Document()
        add_markdown_to_document(doc, "> This is a quote")

        assert any("This is a quote" in p.text for p in doc.paragraphs)

    def test_blockquote_has_left_indent(self) -> None:
        doc = Document()
        add_markdown_to_document(doc, "> Indented text")

        indented = [
            p for p in doc.paragraphs
            if p.paragraph_format.left_indent is not None
            and p.paragraph_format.left_indent > 0
        ]
        assert len(indented) >= 1
        assert indented[0].paragraph_format.left_indent == Inches(0.5)

    def test_multiline_blockquote(self) -> None:
        doc = Document()
        add_markdown_to_document(doc, "> Line one\n> Line two")

        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Line one" in all_text
        assert "Line two" in all_text


# ---------------------------------------------------------------------------
# 8.15 — Link conversion
# ---------------------------------------------------------------------------


class TestLinkConversion:
    """Markdown link converts to DOCX hyperlink."""

    def test_link_creates_hyperlink_element(self) -> None:
        doc = Document()
        add_markdown_to_document(doc, "[Example](http://example.com)")

        # Find the paragraph containing the hyperlink
        for p in doc.paragraphs:
            hyperlinks = p._element.findall(qn("w:hyperlink"))
            if hyperlinks:
                break
        else:
            raise AssertionError("No hyperlink element found")

        assert len(hyperlinks) == 1

    def test_link_text_is_correct(self) -> None:
        doc = Document()
        add_markdown_to_document(doc, "[Click me](http://example.com)")

        for p in doc.paragraphs:
            hyperlinks = p._element.findall(qn("w:hyperlink"))
            if hyperlinks:
                hl = hyperlinks[0]
                runs = hl.findall(qn("w:r"))
                texts = []
                for r in runs:
                    t = r.find(qn("w:t"))
                    if t is not None and t.text:
                        texts.append(t.text)
                assert "Click me" in " ".join(texts)
                return

        raise AssertionError("No hyperlink found")

    def test_link_has_blue_underline_style(self) -> None:
        doc = Document()
        add_markdown_to_document(doc, "[Link](http://example.com)")

        for p in doc.paragraphs:
            hyperlinks = p._element.findall(qn("w:hyperlink"))
            if hyperlinks:
                hl = hyperlinks[0]
                run = hl.find(qn("w:r"))
                rPr = run.find(qn("w:rPr"))
                color = rPr.find(qn("w:color"))
                underline = rPr.find(qn("w:u"))
                assert color is not None
                assert color.get(qn("w:val")) == "0563C1"
                assert underline is not None
                assert underline.get(qn("w:val")) == "single"
                return

        raise AssertionError("No hyperlink found")


# ---------------------------------------------------------------------------
# 8.16 — Nested bullet list indentation
# ---------------------------------------------------------------------------


class TestNestedListIndentation:
    """Nested bullet list produces correct indentation levels."""

    def test_top_level_has_no_extra_indent(self) -> None:
        doc = Document()
        add_markdown_to_document(doc, "- top item")

        list_paras = [
            p for p in doc.paragraphs if "List" in (p.style.name or "")
        ]
        assert len(list_paras) == 1
        assert list_paras[0].paragraph_format.left_indent is None

    def test_nested_item_has_indent(self) -> None:
        doc = Document()
        add_markdown_to_document(doc, "- top\n  - nested")

        list_paras = [
            p for p in doc.paragraphs if "List" in (p.style.name or "")
        ]
        assert len(list_paras) == 2
        # Nested item should have left indent
        nested = list_paras[1]
        assert nested.paragraph_format.left_indent is not None
        assert nested.paragraph_format.left_indent == Inches(0.5)

    def test_double_nested_has_deeper_indent(self) -> None:
        doc = Document()
        add_markdown_to_document(doc, "- top\n  - mid\n    - deep")

        list_paras = [
            p for p in doc.paragraphs if "List" in (p.style.name or "")
        ]
        assert len(list_paras) == 3
        # Deep nested should have 2x indent
        deep = list_paras[2]
        assert deep.paragraph_format.left_indent is not None
        assert deep.paragraph_format.left_indent == Inches(1.0)


# ---------------------------------------------------------------------------
# 8.17 — Existing features still work
# ---------------------------------------------------------------------------


class TestExistingFeatures:
    """Headings, bold, italic, code, and flat lists still work correctly."""

    def test_heading_levels(self) -> None:
        doc = Document()
        add_markdown_to_document(doc, "## H2\n\n### H3\n\n#### H4")

        headings = [
            p for p in doc.paragraphs if p.style.name.startswith("Heading")
        ]
        assert len(headings) == 3
        assert headings[0].text == "H2"
        assert headings[1].text == "H3"
        assert headings[2].text == "H4"

    def test_bold_formatting(self) -> None:
        doc = Document()
        add_markdown_to_document(doc, "**bold text**")

        p = doc.paragraphs[0]
        bold_runs = [r for r in p.runs if r.bold]
        assert any("bold text" in r.text for r in bold_runs)

    def test_italic_formatting(self) -> None:
        doc = Document()
        add_markdown_to_document(doc, "*italic text*")

        p = doc.paragraphs[0]
        italic_runs = [r for r in p.runs if r.italic]
        assert any("italic text" in r.text for r in italic_runs)

    def test_bold_italic_formatting(self) -> None:
        doc = Document()
        add_markdown_to_document(doc, "***bold italic***")

        p = doc.paragraphs[0]
        bi_runs = [r for r in p.runs if r.bold and r.italic]
        assert any("bold italic" in r.text for r in bi_runs)

    def test_inline_code_formatting(self) -> None:
        doc = Document()
        add_markdown_to_document(doc, "`some code`")

        p = doc.paragraphs[0]
        code_runs = [r for r in p.runs if r.font.name == "Courier New"]
        assert any("some code" in r.text for r in code_runs)

    def test_bullet_list(self) -> None:
        doc = Document()
        add_markdown_to_document(doc, "- item A\n- item B\n- item C")

        list_paras = [
            p for p in doc.paragraphs if "Bullet" in (p.style.name or "")
        ]
        assert len(list_paras) == 3
        texts = [p.text for p in list_paras]
        assert "item A" in texts
        assert "item B" in texts
        assert "item C" in texts

    def test_numbered_list(self) -> None:
        doc = Document()
        add_markdown_to_document(doc, "1. first\n2. second\n3. third")

        list_paras = [
            p for p in doc.paragraphs if "Number" in (p.style.name or "")
        ]
        assert len(list_paras) == 3
        texts = [p.text for p in list_paras]
        assert "first" in texts
        assert "second" in texts
        assert "third" in texts

    def test_horizontal_rule(self) -> None:
        doc = Document()
        add_markdown_to_document(doc, "---")

        # Should have a paragraph with a bottom border
        found_hr = False
        for p in doc.paragraphs:
            pPr = p._element.find(qn("w:pPr"))
            if pPr is not None:
                pBdr = pPr.find(qn("w:pBdr"))
                if pBdr is not None:
                    bottom = pBdr.find(qn("w:bottom"))
                    if bottom is not None:
                        found_hr = True
                        break
        assert found_hr, "No horizontal rule (bottom border) found"

    def test_empty_input_produces_no_content(self) -> None:
        doc = Document()
        add_markdown_to_document(doc, "")
        assert len(doc.paragraphs) == 0

    def test_none_like_whitespace_input(self) -> None:
        doc = Document()
        add_markdown_to_document(doc, "   \n\n  ")
        assert len(doc.paragraphs) == 0
