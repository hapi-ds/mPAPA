"""Unit tests for the DOCX exporter module."""

import pytest
from pathlib import Path

from docx import Document

from patent_system.export.docx_exporter import DOCXExporter, validate_export


# --- validate_export ---


class TestValidateExport:
    """Tests for the standalone validate_export function."""

    def test_both_non_empty_returns_true(self) -> None:
        assert validate_export("claim 1", "description text") is True

    def test_claims_none_returns_false(self) -> None:
        assert validate_export(None, "description text") is False

    def test_description_none_returns_false(self) -> None:
        assert validate_export("claim 1", None) is False

    def test_both_none_returns_false(self) -> None:
        assert validate_export(None, None) is False

    def test_claims_empty_string_returns_false(self) -> None:
        assert validate_export("", "description text") is False

    def test_description_empty_string_returns_false(self) -> None:
        assert validate_export("claim 1", "") is False

    def test_both_empty_strings_returns_false(self) -> None:
        assert validate_export("", "") is False


# --- DOCXExporter ---


class TestDOCXExporterResolveTemplate:
    """Tests for _resolve_template."""

    def test_none_template_name_returns_none(self, tmp_path: Path) -> None:
        exporter = DOCXExporter(template_dir=tmp_path, template_name=None)
        assert exporter._resolve_template() is None

    def test_missing_file_returns_none(self, tmp_path: Path) -> None:
        exporter = DOCXExporter(template_dir=tmp_path, template_name="missing.docx")
        assert exporter._resolve_template() is None

    def test_existing_file_returns_path(self, tmp_path: Path) -> None:
        # Create a valid .docx template
        template_path = tmp_path / "my_template.docx"
        Document().save(str(template_path))

        exporter = DOCXExporter(template_dir=tmp_path, template_name="my_template.docx")
        result = exporter._resolve_template()
        assert result == template_path


class TestDOCXExporterExport:
    """Tests for the export method."""

    def test_export_blank_document(self, tmp_path: Path) -> None:
        exporter = DOCXExporter(template_dir=tmp_path, template_name=None)
        output = tmp_path / "output.docx"

        result = exporter.export("My claims", "My description", output)

        assert result == output
        assert output.exists()

        doc = Document(str(output))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "My claims" in full_text
        assert "My description" in full_text

    def test_export_with_template(self, tmp_path: Path) -> None:
        # Create a template with a pre-existing paragraph
        template_path = tmp_path / "template.docx"
        tmpl = Document()
        tmpl.add_paragraph("Template header content")
        tmpl.save(str(template_path))

        exporter = DOCXExporter(template_dir=tmp_path, template_name="template.docx")
        output = tmp_path / "output.docx"

        result = exporter.export("Claim A", "Description B", output)

        assert result == output
        doc = Document(str(output))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Template header content" in full_text
        assert "Claim A" in full_text
        assert "Description B" in full_text

    def test_export_creates_parent_directories(self, tmp_path: Path) -> None:
        exporter = DOCXExporter(template_dir=tmp_path, template_name=None)
        output = tmp_path / "nested" / "dir" / "output.docx"

        exporter.export("claims", "description", output)
        assert output.exists()

    def test_export_contains_headings(self, tmp_path: Path) -> None:
        exporter = DOCXExporter(template_dir=tmp_path, template_name=None)
        output = tmp_path / "output.docx"

        exporter.export("claims text", "description text", output)

        doc = Document(str(output))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Claims" in headings
        assert "Description" in headings

    def test_export_falls_back_when_template_missing(self, tmp_path: Path) -> None:
        exporter = DOCXExporter(template_dir=tmp_path, template_name="nonexistent.docx")
        output = tmp_path / "output.docx"

        result = exporter.export("claims", "description", output)

        assert result == output
        assert output.exists()
        doc = Document(str(output))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "claims" in full_text
        assert "description" in full_text


class TestDOCXExporterListTemplates:
    """Tests for list_available_templates."""

    def test_empty_directory(self, tmp_path: Path) -> None:
        exporter = DOCXExporter(template_dir=tmp_path)
        assert exporter.list_available_templates() == []

    def test_nonexistent_directory(self, tmp_path: Path) -> None:
        exporter = DOCXExporter(template_dir=tmp_path / "nonexistent")
        assert exporter.list_available_templates() == []

    def test_lists_only_docx_files(self, tmp_path: Path) -> None:
        # Create a mix of files
        Document().save(str(tmp_path / "a.docx"))
        Document().save(str(tmp_path / "b.docx"))
        (tmp_path / "readme.txt").write_text("not a template")
        (tmp_path / ".gitkeep").write_text("")

        exporter = DOCXExporter(template_dir=tmp_path)
        result = exporter.list_available_templates()
        assert result == ["a.docx", "b.docx"]

    def test_returns_sorted_list(self, tmp_path: Path) -> None:
        Document().save(str(tmp_path / "z_template.docx"))
        Document().save(str(tmp_path / "a_template.docx"))

        exporter = DOCXExporter(template_dir=tmp_path)
        result = exporter.list_available_templates()
        assert result == ["a_template.docx", "z_template.docx"]


class TestDOCXExporterWorkflowSteps:
    """Tests for the workflow_steps parameter in export."""

    def test_export_without_workflow_steps(self, tmp_path: Path) -> None:
        """Existing behavior is unchanged when workflow_steps is None."""
        exporter = DOCXExporter(template_dir=tmp_path, template_name=None)
        output = tmp_path / "output.docx"

        exporter.export("claims", "description", output, workflow_steps=None)

        doc = Document(str(output))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert headings == ["Claims", "Description"]

    def test_export_with_empty_workflow_steps_dict(self, tmp_path: Path) -> None:
        """Empty dict produces no extra sections."""
        exporter = DOCXExporter(template_dir=tmp_path, template_name=None)
        output = tmp_path / "output.docx"

        exporter.export("claims", "description", output, workflow_steps={})

        doc = Document(str(output))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert headings == ["Claims", "Description"]

    def test_export_with_workflow_steps_adds_sections(self, tmp_path: Path) -> None:
        """Non-empty workflow steps appear as heading-1 sections with display names."""
        exporter = DOCXExporter(template_dir=tmp_path, template_name=None)
        output = tmp_path / "output.docx"

        steps = {
            "initial_idea": "My invention idea",
            "novelty_analysis": "Novelty looks good",
        }
        exporter.export("claims", "description", output, workflow_steps=steps)

        doc = Document(str(output))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert headings == ["Claims", "Description", "Initial Idea", "Novelty Analysis"]

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "My invention idea" in full_text
        assert "Novelty looks good" in full_text

    def test_export_omits_empty_content_steps(self, tmp_path: Path) -> None:
        """Steps with empty or whitespace-only content are omitted."""
        exporter = DOCXExporter(template_dir=tmp_path, template_name=None)
        output = tmp_path / "output.docx"

        steps = {
            "initial_idea": "Has content",
            "claims_drafting": "",
            "prior_art_search": "   ",
            "novelty_analysis": "Also has content",
        }
        exporter.export("claims", "description", output, workflow_steps=steps)

        doc = Document(str(output))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert headings == ["Claims", "Description", "Initial Idea", "Novelty Analysis"]

    def test_export_workflow_steps_in_canonical_order(self, tmp_path: Path) -> None:
        """Steps appear in canonical order regardless of dict insertion order."""
        exporter = DOCXExporter(template_dir=tmp_path, template_name=None)
        output = tmp_path / "output.docx"

        # Insert in reverse order
        steps = {
            "patent_draft": "Draft content",
            "market_potential": "Market content",
            "initial_idea": "Idea content",
        }
        exporter.export("claims", "description", output, workflow_steps=steps)

        doc = Document(str(output))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert headings == [
            "Claims", "Description",
            "Initial Idea", "Market Potential",
        ]

    def test_export_workflow_steps_before_references_and_chat(self, tmp_path: Path) -> None:
        """Workflow steps appear after Description but before References and AI Chat Log."""
        exporter = DOCXExporter(template_dir=tmp_path, template_name=None)
        output = tmp_path / "output.docx"

        steps = {"consistency_review": "Review notes"}
        refs = [{"title": "Ref 1", "abstract": "Abstract 1"}]
        chat = [{"role": "user", "message": "Hello"}]

        exporter.export(
            "claims", "description", output,
            references=refs, chat_history=chat, workflow_steps=steps,
        )

        doc = Document(str(output))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        # Heading 1 sections only
        h1_headings = [
            p.text for p in doc.paragraphs if p.style.name == "Heading 1"
        ]
        assert h1_headings == [
            "Claims", "Description", "Consistency Review",
            "References", "AI Chat Log",
        ]

    def test_export_all_workflow_steps(self, tmp_path: Path) -> None:
        """All nine steps with content appear in canonical order."""
        from patent_system.export.docx_exporter import WORKFLOW_STEP_ORDER, STEP_DISPLAY_NAMES

        exporter = DOCXExporter(template_dir=tmp_path, template_name=None)
        output = tmp_path / "output.docx"

        steps = {key: f"Content for {key}" for key in WORKFLOW_STEP_ORDER}
        exporter.export("claims", "description", output, workflow_steps=steps)

        doc = Document(str(output))
        h1_headings = [
            p.text for p in doc.paragraphs if p.style.name == "Heading 1"
        ]
        expected = ["Claims", "Description"] + [
            STEP_DISPLAY_NAMES[k] for k in WORKFLOW_STEP_ORDER
            if k != "patent_draft"
        ]
        assert h1_headings == expected
