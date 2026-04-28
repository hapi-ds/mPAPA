"""Property-based tests for the markdown-to-DOCX converter.

**Validates: Requirements 12.6**

Property 7: Markdown-to-DOCX round-trip preserves text content.
For any markdown string containing headings, bold, italic, and plain text,
the DOCX output SHALL contain all the plain text content from the markdown
input (formatting may differ but text content is preserved).
"""

import re

from docx import Document
from hypothesis import given, settings
from hypothesis import strategies as st

from patent_system.export.markdown_converter import add_markdown_to_document

# ---------------------------------------------------------------------------
# Strategies
# ---------------------------------------------------------------------------

# XML-safe printable characters (python-docx/lxml rejects NULL bytes and
# certain control characters).
_xml_safe_chars = st.characters(
    categories=("L", "N", "Z"),
    exclude_characters="\x00\r",
)

# Short non-empty words for building markdown fragments.
_word = st.text(alphabet=_xml_safe_chars, min_size=1, max_size=20).filter(
    lambda s: s.strip() != ""
)


def _plain_text(word: str) -> str:
    """Wrap a word as plain text."""
    return word


def _bold_text(word: str) -> str:
    """Wrap a word as bold markdown."""
    return f"**{word}**"


def _italic_text(word: str) -> str:
    """Wrap a word as italic markdown."""
    return f"*{word}*"


def _heading_text(word: str) -> str:
    """Wrap a word as a level-2 heading."""
    return f"## {word}"


# Strategy that generates a markdown document with a mix of elements.
_md_fragment = st.one_of(
    _word.map(_plain_text),
    _word.map(_bold_text),
    _word.map(_italic_text),
    _word.map(_heading_text),
)

_md_document = st.lists(_md_fragment, min_size=1, max_size=8).map(
    lambda parts: "\n\n".join(parts)
)


def _strip_markdown(md_text: str) -> list[str]:
    """Extract plain text words from markdown, stripping formatting markers."""
    # Remove heading markers
    text = re.sub(r"^#{1,6}\s+", "", md_text, flags=re.MULTILINE)
    # Remove bold/italic markers
    text = text.replace("***", "").replace("**", "").replace("*", "")
    # Remove inline code markers
    text = text.replace("`", "")
    # Split into words and filter empties
    words = text.split()
    return [w for w in words if w.strip()]


def _extract_docx_text(doc: Document) -> str:
    """Extract all text from a DOCX document (paragraphs + tables)."""
    parts: list[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return " ".join(parts)


# ---------------------------------------------------------------------------
# Property 7: Markdown-to-DOCX preserves all plain text content
# ---------------------------------------------------------------------------


class TestMarkdownToDocxPreservesText:
    """Property 7: Markdown-to-DOCX round-trip preserves text content.

    For any markdown string containing headings, bold, italic, and plain text,
    the DOCX output SHALL contain all the plain text content from the markdown
    input.

    **Validates: Requirements 12.6**
    """

    @given(md_text=_md_document)
    @settings(max_examples=100)
    def test_all_plain_text_preserved_in_docx(self, md_text: str) -> None:
        """All plain text words from the markdown appear in the DOCX output."""
        doc = Document()
        add_markdown_to_document(doc, md_text)

        docx_text = _extract_docx_text(doc)
        plain_words = _strip_markdown(md_text)

        for word in plain_words:
            assert word in docx_text, (
                f"Word {word!r} from markdown not found in DOCX output.\n"
                f"Markdown: {md_text!r}\n"
                f"DOCX text: {docx_text!r}"
            )
