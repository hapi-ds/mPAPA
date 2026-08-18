"""Report rendering module for patent review findings.

Renders a ReviewReport to Markdown string or DOCX file.
"""

from __future__ import annotations

from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from .engine import ReviewFinding, ReviewReport

# Severity ordering for grouping (critical first)
_SEVERITY_ORDER: list[str] = ["critical", "major", "minor", "observation"]

_SEVERITY_BADGES: dict[str, str] = {
    "critical": "🔴 Critical",
    "major": "🟠 Major",
    "minor": "🟡 Minor",
    "observation": "🔵 Observation",
}

_SEVERITY_COLORS: dict[str, RGBColor] = {
    "critical": RGBColor(0xCC, 0x00, 0x00),
    "major": RGBColor(0xE6, 0x7E, 0x00),
    "minor": RGBColor(0xCC, 0xAA, 0x00),
    "observation": RGBColor(0x33, 0x66, 0xCC),
}


def _count_by_severity(report: ReviewReport) -> Counter[str]:
    """Count findings grouped by severity level."""
    counts: Counter[str] = Counter()
    for finding in report.findings:
        counts[finding.severity.lower()] += 1
    return counts


def _group_by_severity(findings: list[ReviewFinding]) -> dict[str, list[ReviewFinding]]:
    """Group findings by severity, ordered critical-first."""
    grouped: dict[str, list[ReviewFinding]] = {s: [] for s in _SEVERITY_ORDER}
    for finding in findings:
        key = finding.severity.lower()
        if key in grouped:
            grouped[key].append(finding)
        else:
            grouped.setdefault(key, []).append(finding)
    return grouped


def render_markdown(report: ReviewReport) -> str:
    """Render a ReviewReport as a Markdown string.

    Produces a summary header with severity statistics, followed by each
    finding as a section with rule_id, severity badge, location, finding
    text, suggestion, and reference.

    Args:
        report: The review report to render.

    Returns:
        Formatted Markdown string.
    """
    lines: list[str] = []

    # Title
    title = f"{report.jurisdiction} Patent Review"
    lines.append(f"# Review Report: {title}")
    lines.append("")

    # Summary statistics
    counts = _count_by_severity(report)
    stats_parts: list[str] = []
    for severity in _SEVERITY_ORDER:
        count = counts.get(severity, 0)
        badge = _SEVERITY_BADGES.get(severity, severity.title())
        stats_parts.append(f"**{count}** {badge}")

    lines.append("## Summary")
    lines.append("")
    lines.append(" | ".join(stats_parts))
    lines.append("")
    lines.append(f"**Total findings:** {len(report.findings)}")
    lines.append("")

    # Findings
    lines.append("## Findings")
    lines.append("")

    for i, finding in enumerate(report.findings, start=1):
        severity_badge = _SEVERITY_BADGES.get(
            finding.severity.lower(), finding.severity
        )
        lines.append(f"### {i}. [{finding.rule_id}] {severity_badge}")
        lines.append("")

        if finding.location:
            lines.append(f"**Location:** `{finding.location}`")
            lines.append("")

        lines.append(f"**Finding:** {finding.finding}")
        lines.append("")

        if finding.suggestion:
            lines.append(f"**Suggestion:** {finding.suggestion}")
            lines.append("")

        if finding.reference:
            lines.append(f"**Reference:** {finding.reference}")
            lines.append("")

        lines.append("---")
        lines.append("")

    return "\n".join(lines)


def render_docx(report: ReviewReport, output_path: Path) -> Path:
    """Render a ReviewReport as a professional DOCX document.

    Produces a document with a title page, summary table, and findings
    grouped by severity (critical first).

    Args:
        report: The review report to render.
        output_path: Path where the .docx file will be written.

    Returns:
        The output_path after successful write.
    """
    doc = Document()

    # -- Title page --
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(f"Review Report\n{report.jurisdiction} Patent Review")
    title_run.font.size = Pt(26)
    title_run.bold = True

    doc.add_paragraph()  # spacer

    if hasattr(report, "date") and report.date:
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(f"Date: {report.date}").font.size = Pt(12)

    doc.add_page_break()

    # -- Summary table --
    doc.add_heading("Summary", level=1)

    counts = _count_by_severity(report)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Severity"
    hdr_cells[1].text = "Count"
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Data rows
    for severity in _SEVERITY_ORDER:
        count = counts.get(severity, 0)
        row_cells = table.add_row().cells
        row_cells[0].text = severity.title()
        row_cells[1].text = str(count)

        # Color the severity cell text
        color = _SEVERITY_COLORS.get(severity)
        if color:
            for paragraph in row_cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = color

    # Total row
    total_cells = table.add_row().cells
    total_cells[0].text = "Total"
    total_cells[1].text = str(len(report.findings))
    for cell in [total_cells[0], total_cells[1]]:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    doc.add_paragraph()  # spacer

    # -- Findings grouped by severity --
    doc.add_heading("Findings", level=1)

    grouped = _group_by_severity(report.findings)
    finding_num = 0

    for severity in _SEVERITY_ORDER:
        findings_in_group = grouped.get(severity, [])
        if not findings_in_group:
            continue

        doc.add_heading(f"{severity.title()} ({len(findings_in_group)})", level=2)

        for finding in findings_in_group:
            finding_num += 1

            # Finding heading
            heading = doc.add_heading(level=3)
            heading_run = heading.add_run(
                f"{finding_num}. [{finding.rule_id}] — {severity.title()}"
            )
            color = _SEVERITY_COLORS.get(severity)
            if color:
                heading_run.font.color.rgb = color

            # Location
            if finding.location:
                loc_para = doc.add_paragraph()
                loc_para.add_run("Location: ").bold = True
                loc_para.add_run(finding.location)

            # Finding text
            finding_para = doc.add_paragraph()
            finding_para.add_run("Finding: ").bold = True
            finding_para.add_run(finding.finding)

            # Suggestion
            if finding.suggestion:
                sug_para = doc.add_paragraph()
                sug_para.add_run("Suggestion: ").bold = True
                sug_para.add_run(finding.suggestion)

            # Reference
            if finding.reference:
                ref_para = doc.add_paragraph()
                ref_para.add_run("Reference: ").bold = True
                ref_para.add_run(finding.reference)

            # Separator
            doc.add_paragraph("─" * 60)

    # Write file
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
