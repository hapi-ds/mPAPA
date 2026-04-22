"""DOCX generation with optional template support.

Generates .docx patent documents containing claims and description sections.
Supports user-provided .docx templates for custom formatting; falls back to
a blank Document when no template is configured or the file is missing.
"""

import logging
from pathlib import Path

from docx import Document
from docx.shared import Pt

from patent_system.export.markdown_converter import add_markdown_to_document

logger = logging.getLogger(__name__)

WORKFLOW_STEP_ORDER: list[str] = [
    "initial_idea",
    "claims_drafting",
    "prior_art_search",
    "novelty_analysis",
    "consistency_review",
    "market_potential",
    "legal_clarification",
    "disclosure_summary",
    "patent_draft",
]

STEP_DISPLAY_NAMES: dict[str, str] = {
    "initial_idea": "Initial Idea",
    "claims_drafting": "Claims Drafting",
    "prior_art_search": "Prior Art Search",
    "novelty_analysis": "Novelty Analysis",
    "consistency_review": "Consistency Review",
    "market_potential": "Market Potential",
    "legal_clarification": "Legal Clarification",
    "disclosure_summary": "Disclosure Summary",
    "patent_draft": "Patent Draft",
}


def _safe_add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a heading, falling back to a bold paragraph if the style is missing."""
    try:
        doc.add_heading(text, level=level)
    except KeyError:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        size_map = {1: Pt(18), 2: Pt(16), 3: Pt(14), 4: Pt(12)}
        run.font.size = size_map.get(level, Pt(14))


def _safe_add_list_paragraph(doc: Document, text: str, style: str) -> None:
    """Add a list paragraph, falling back to a prefixed normal paragraph."""
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        p = doc.add_paragraph()
        prefix = "• " if "Bullet" in style else ""
        if prefix:
            p.add_run(prefix)
    p.add_run(text)


def validate_export(claims: str | None, description: str | None) -> bool:
    """Check whether claims and description are non-empty and suitable for export.

    Args:
        claims: The patent claims text.
        description: The patent description text.

    Returns:
        True if both claims and description are non-empty strings, False otherwise.
    """
    if not claims or not description:
        return False
    return True


class DOCXExporter:
    """Generates .docx patent documents using optional templates from the templates directory."""

    def __init__(self, template_dir: Path, template_name: str | None = None) -> None:
        """Initialize the exporter.

        Args:
            template_dir: Path to the templates directory containing .docx template files.
            template_name: Filename of the template to use (e.g. "european_patent.docx").
                           If None or file not found, starts with a blank Document().
        """
        self._template_dir = template_dir
        self._template_name = template_name

    def _resolve_template(self) -> Path | None:
        """Resolve the template path from template_dir + template_name.

        Returns:
            The full Path to the template file, or None if template_name is None
            or the file does not exist.
        """
        if self._template_name is None:
            return None

        template_path = self._template_dir / self._template_name
        if not template_path.exists():
            logger.warning(
                "Configured template '%s' not found in '%s'; falling back to blank document.",
                self._template_name,
                self._template_dir,
            )
            return None

        return template_path

    def export(
        self,
        claims: str,
        description: str,
        output_path: Path,
        references: list[dict] | None = None,
        chat_history: list[dict] | None = None,
        workflow_steps: dict[str, str] | None = None,
    ) -> Path:
        """Generate a .docx file with claims, description, references, and chat.

        Args:
            claims: The patent claims text.
            description: The patent description text.
            output_path: Destination path for the generated .docx file.
            references: Optional list of dicts with 'title', 'abstract',
                'source', and optionally 'patent_number' or 'doi' keys.
            chat_history: Optional list of dicts with 'role' and 'message' keys.
            workflow_steps: Optional mapping of step_key to content text.
                Non-empty steps are included as heading-1 sections in canonical
                order between Description and References.

        Returns:
            The output_path where the document was saved.
        """
        template_path = self._resolve_template()

        if template_path is not None:
            doc = Document(str(template_path))
        else:
            doc = Document()

        _safe_add_heading(doc, "Claims", level=1)
        add_markdown_to_document(doc, claims)

        _safe_add_heading(doc, "Description", level=1)
        add_markdown_to_document(doc, description)

        if workflow_steps:
            for step_key in WORKFLOW_STEP_ORDER:
                if step_key == "patent_draft":
                    continue
                content = workflow_steps.get(step_key, "")
                if content and content.strip():
                    display_name = STEP_DISPLAY_NAMES.get(step_key, step_key)
                    _safe_add_heading(doc, display_name, level=1)
                    add_markdown_to_document(doc, content)

        if references:
            _safe_add_heading(doc, "References", level=1)
            for i, ref in enumerate(references, 1):
                title = ref.get("title", "Untitled")
                source = ref.get("source", "")
                record_id = ref.get("patent_number") or ref.get("doi") or ""
                abstract = ref.get("abstract", "")
                has_full_text = ref.get("has_full_text", False)
                url = ref.get("url", "")
                relevance = ref.get("relevance_score")

                # Heading: [1] Title (Source)
                heading_text = f"[{i}] {title}"
                if source:
                    heading_text += f" ({source})"
                _safe_add_heading(doc, heading_text, level=2)

                # Metadata line: ID, relevance, full text indicator
                meta_parts: list[str] = []
                if record_id and record_id != "UNKNOWN":
                    meta_parts.append(record_id)
                if relevance is not None:
                    meta_parts.append(f"Relevance: {relevance}%")
                meta_parts.append("Full Text" if has_full_text else "Abstract Only")
                if meta_parts:
                    doc.add_paragraph(" · ".join(meta_parts)).runs[0].italic = True

                # Clickable link
                if url:
                    p = doc.add_paragraph()
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement
                    hyperlink = OxmlElement("w:hyperlink")
                    hyperlink.set(qn("r:id"), p.part.relate_to(
                        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True,
                    ))
                    run_el = OxmlElement("w:r")
                    rPr = OxmlElement("w:rPr")
                    color = OxmlElement("w:color")
                    color.set(qn("w:val"), "0563C1")
                    rPr.append(color)
                    u = OxmlElement("w:u")
                    u.set(qn("w:val"), "single")
                    rPr.append(u)
                    run_el.append(rPr)
                    text_el = OxmlElement("w:t")
                    text_el.text = url
                    run_el.append(text_el)
                    hyperlink.append(run_el)
                    p._element.append(hyperlink)

                if abstract:
                    doc.add_paragraph(abstract)

        if chat_history:
            _safe_add_heading(doc, "AI Chat Log", level=1)
            for msg in chat_history:
                role = msg.get("role", "unknown")
                text = msg.get("message", "")
                label = "You" if role == "user" else "Assistant"
                p = doc.add_paragraph()
                run = p.add_run(f"{label}: ")
                run.bold = True
                p.add_run(text)

        # Ensure parent directories exist
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        return output_path

    def list_available_templates(self) -> list[str]:
        """List all .docx files in the templates directory.

        Returns:
            A list of .docx filenames found in the template directory.
            Returns an empty list if the directory does not exist.
        """
        if not self._template_dir.exists():
            return []

        return sorted(f.name for f in self._template_dir.iterdir() if f.suffix == ".docx")
