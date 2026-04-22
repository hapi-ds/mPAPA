"""Markdown-to-DOCX converter using mistune AST parsing.

Parses markdown text into an AST via mistune, then walks the tree to produce
python-docx elements. Replaces the regex-based ``_add_markdown_content``
approach with proper parsing for headings, inline formatting, lists (including
nested), tables, blockquotes, horizontal rules, and hyperlinks.
"""

from __future__ import annotations

import logging
from typing import Any

import mistune
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Mistune markdown parser (AST mode with table plugin)
# ---------------------------------------------------------------------------
_md_parser = mistune.create_markdown(renderer="ast", plugins=["table"])


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def add_markdown_to_document(doc: Document, markdown_text: str) -> None:
    """Parse *markdown_text* and append formatted content to *doc*.

    Args:
        doc: A ``python-docx`` Document instance to append to.
        markdown_text: Markdown-formatted string to convert.
    """
    if not markdown_text or not markdown_text.strip():
        return

    tokens: list[dict[str, Any]] = _md_parser(markdown_text)
    _walk_tokens(doc, tokens, list_depth=0)


# ---------------------------------------------------------------------------
# AST walker
# ---------------------------------------------------------------------------


def _walk_tokens(
    doc: Document,
    tokens: list[dict[str, Any]],
    *,
    list_depth: int = 0,
) -> None:
    """Recursively walk AST tokens and add DOCX elements to *doc*."""
    for token in tokens:
        tok_type = token.get("type", "")

        if tok_type == "heading":
            _handle_heading(doc, token)

        elif tok_type == "paragraph":
            _handle_paragraph(doc, token)

        elif tok_type == "list":
            _handle_list(doc, token, list_depth=list_depth)

        elif tok_type == "table":
            _handle_table(doc, token)

        elif tok_type == "block_quote":
            _handle_blockquote(doc, token)

        elif tok_type == "thematic_break":
            _handle_thematic_break(doc)

        elif tok_type == "block_code":
            _handle_block_code(doc, token)

        elif tok_type == "blank_line":
            # Skip blank line tokens — they are structural separators
            continue

        else:
            # Fallback: if the token has children, recurse
            children = token.get("children")
            if children and isinstance(children, list):
                _walk_tokens(doc, children, list_depth=list_depth)


# ---------------------------------------------------------------------------
# Block-level handlers
# ---------------------------------------------------------------------------


def _handle_heading(doc: Document, token: dict[str, Any]) -> None:
    """Convert a heading token to a DOCX heading (levels 2-6 → DOCX 2-4)."""
    level = token.get("attrs", {}).get("level", 2)
    # Clamp to DOCX heading levels 1-4
    docx_level = min(max(level, 1), 4)

    try:
        heading = doc.add_heading(level=docx_level)
    except KeyError:
        heading = doc.add_paragraph()
        run = heading.add_run()
        run.bold = True
        size_map = {1: Pt(18), 2: Pt(16), 3: Pt(14), 4: Pt(12)}
        run.font.size = size_map.get(docx_level, Pt(14))

    _add_inline_children(heading, token.get("children", []))


def _handle_paragraph(doc: Document, token: dict[str, Any]) -> None:
    """Convert a paragraph token to a DOCX paragraph with inline formatting."""
    p = doc.add_paragraph()
    _add_inline_children(p, token.get("children", []))


def _handle_list(
    doc: Document,
    token: dict[str, Any],
    *,
    list_depth: int = 0,
) -> None:
    """Convert a list token (bullet or numbered) to DOCX list paragraphs."""
    attrs = token.get("attrs", {})
    ordered = attrs.get("ordered", False)

    for item in token.get("children", []):
        if item.get("type") != "list_item":
            continue
        for child in item.get("children", []):
            child_type = child.get("type", "")
            if child_type == "list":
                # Nested list — increase depth
                _handle_list(doc, child, list_depth=list_depth + 1)
            elif child_type in ("block_text", "paragraph"):
                _add_list_paragraph(
                    doc,
                    child.get("children", []),
                    ordered=ordered,
                    depth=list_depth,
                )


def _add_list_paragraph(
    doc: Document,
    children: list[dict[str, Any]],
    *,
    ordered: bool,
    depth: int,
) -> None:
    """Add a single list item paragraph with appropriate style and indent."""
    style_name = "List Number" if ordered else "List Bullet"
    try:
        p = doc.add_paragraph(style=style_name)
    except KeyError:
        p = doc.add_paragraph()
        if not ordered:
            p.add_run("• ")

    # Apply indentation for nested lists
    if depth > 0:
        p.paragraph_format.left_indent = Inches(0.5 * depth)

    _add_inline_children(p, children)


def _handle_table(doc: Document, token: dict[str, Any]) -> None:
    """Convert a table token to a DOCX table with cell borders."""
    # Collect header and body rows
    header_cells: list[list[dict[str, Any]]] = []
    body_rows: list[list[list[dict[str, Any]]]] = []

    for child in token.get("children", []):
        child_type = child.get("type", "")
        if child_type == "table_head":
            header_cells = [
                cell.get("children", [])
                for cell in child.get("children", [])
                if cell.get("type") == "table_cell"
            ]
        elif child_type == "table_body":
            for row in child.get("children", []):
                if row.get("type") == "table_row":
                    row_cells = [
                        cell.get("children", [])
                        for cell in row.get("children", [])
                        if cell.get("type") == "table_cell"
                    ]
                    body_rows.append(row_cells)

    num_cols = len(header_cells) if header_cells else (
        len(body_rows[0]) if body_rows else 0
    )
    num_rows = (1 if header_cells else 0) + len(body_rows)

    if num_cols == 0 or num_rows == 0:
        return

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = "Table Grid"

    # Fill header row
    if header_cells:
        for col_idx, cell_children in enumerate(header_cells):
            if col_idx < num_cols:
                cell = table.rows[0].cells[col_idx]
                cell.paragraphs[0].clear()
                run = cell.paragraphs[0].add_run()
                run.bold = True
                _add_inline_children(cell.paragraphs[0], cell_children)
                # Make header runs bold
                for r in cell.paragraphs[0].runs:
                    r.bold = True

    # Fill body rows
    row_offset = 1 if header_cells else 0
    for row_idx, row_cells in enumerate(body_rows):
        for col_idx, cell_children in enumerate(row_cells):
            if col_idx < num_cols and (row_idx + row_offset) < num_rows:
                cell = table.rows[row_idx + row_offset].cells[col_idx]
                cell.paragraphs[0].clear()
                _add_inline_children(cell.paragraphs[0], cell_children)

    # Ensure all cells have borders
    _set_table_borders(table)


def _set_table_borders(table: Any) -> None:
    """Apply single-line borders to all cells in a DOCX table."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    tbl_pr.append(borders)
    if tbl.tblPr is None:
        tbl.append(tbl_pr)


def _handle_blockquote(doc: Document, token: dict[str, Any]) -> None:
    """Convert a blockquote token to an indented DOCX paragraph."""
    children = token.get("children", [])
    for child in children:
        child_type = child.get("type", "")
        if child_type == "paragraph":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            _add_inline_children(p, child.get("children", []))
        elif child_type == "block_quote":
            # Nested blockquote — just recurse (will get its own indent)
            _handle_blockquote(doc, child)
        else:
            # Fallback: create indented paragraph for any other block content
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            inline_children = child.get("children", [])
            if inline_children:
                _add_inline_children(p, inline_children)
            elif child.get("raw"):
                p.add_run(child["raw"])


def _handle_thematic_break(doc: Document) -> None:
    """Convert a thematic break (``---``) to a DOCX horizontal line."""
    p = doc.add_paragraph()
    p_pr = p._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _handle_block_code(doc: Document, token: dict[str, Any]) -> None:
    """Convert a fenced code block to a DOCX paragraph with monospace font."""
    raw = token.get("raw", "")
    p = doc.add_paragraph()
    run = p.add_run(raw.rstrip("\n"))
    run.font.name = "Courier New"
    run.font.size = Pt(9)


# ---------------------------------------------------------------------------
# Inline-level handler
# ---------------------------------------------------------------------------


def _add_inline_children(
    paragraph: Any,
    children: list[dict[str, Any]],
    *,
    bold: bool = False,
    italic: bool = False,
) -> None:
    """Recursively add inline children to a DOCX paragraph.

    Tracks inherited bold/italic state so nested ``***bold italic***``
    renders correctly.
    """
    for child in children:
        child_type = child.get("type", "")

        if child_type == "text":
            run = paragraph.add_run(child.get("raw", ""))
            if bold:
                run.bold = True
            if italic:
                run.italic = True

        elif child_type == "strong":
            _add_inline_children(
                paragraph,
                child.get("children", []),
                bold=True,
                italic=italic,
            )

        elif child_type == "emphasis":
            _add_inline_children(
                paragraph,
                child.get("children", []),
                bold=bold,
                italic=True,
            )

        elif child_type == "codespan":
            run = paragraph.add_run(child.get("raw", ""))
            run.font.name = "Courier New"
            run.font.size = Pt(9)

        elif child_type == "link":
            url = child.get("attrs", {}).get("url", "")
            _add_hyperlink(paragraph, url, child.get("children", []))

        elif child_type == "softbreak":
            paragraph.add_run(" ")

        elif child_type == "linebreak":
            run = paragraph.add_run()
            run.add_break()

        else:
            # Fallback: render raw text if present, or recurse children
            raw = child.get("raw")
            if raw:
                run = paragraph.add_run(raw)
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True
            sub_children = child.get("children")
            if sub_children and isinstance(sub_children, list):
                _add_inline_children(
                    paragraph, sub_children, bold=bold, italic=italic
                )


def _add_hyperlink(
    paragraph: Any,
    url: str,
    children: list[dict[str, Any]],
) -> None:
    """Add a clickable hyperlink to a DOCX paragraph.

    Uses the low-level OOXML approach to create a ``w:hyperlink`` element
    with blue underlined text.
    """
    # Collect link text from children
    link_text = _extract_text_from_children(children)

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_el = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Blue color
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    # Underline
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    run_el.append(rPr)

    text_el = OxmlElement("w:t")
    text_el.text = link_text
    text_el.set(qn("xml:space"), "preserve")
    run_el.append(text_el)

    hyperlink.append(run_el)
    paragraph._element.append(hyperlink)


def _extract_text_from_children(children: list[dict[str, Any]]) -> str:
    """Recursively extract plain text from inline AST children."""
    parts: list[str] = []
    for child in children:
        raw = child.get("raw")
        if raw:
            parts.append(raw)
        sub = child.get("children")
        if sub and isinstance(sub, list):
            parts.append(_extract_text_from_children(sub))
    return "".join(parts)
